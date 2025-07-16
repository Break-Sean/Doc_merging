import os
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

def merge_word_documents(directory, output_filename):
    # 创建新文档
    merged_doc = Document()
    
    # 设置一级标题样式（可选）
    style = merged_doc.styles['Heading 1']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 遍历目录中的所有文件
    for filename in sorted(os.listdir(directory)):
        if filename.endswith('.docx'):
            filepath = os.path.join(directory, filename)
            print(f"处理文件中: {filename}")
            
            try:
                doc = Document(filepath)
                found_heading = False
                
                # 查找第一个标题作为一级标题
                heading_text = os.path.splitext(filename)[0]  # 默认使用文件名
                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        heading_text = para.text
                        found_heading = True
                        break
                
                # 添加一级标题
                merged_doc.add_heading(heading_text, level=1)
                
                # 复制文档内容
                for element in doc.element.body:
                    # 跳过第一个标题（如果找到）
                    if found_heading and element == doc.paragraphs[0]._element:
                        continue
                    merged_doc.element.body.append(element)
                
                # 添加分页符
                merged_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                
            except Exception as e:
                print(f"处理 {filename} 时出错: {str(e)}")
    
    # 删除最后一个分页符
    if len(merged_doc.paragraphs) > 0:
        last_para = merged_doc.paragraphs[-1]
        if last_para.runs and last_para.runs[-1].text == "":
            last_para._element.getparent().remove(last_para._element)
    
    # 保存合并后的文档
    merged_doc.save(os.path.join(directory, output_filename))
    print(f"合并完成! 输出文件: {output_filename}")

if __name__ == "__main__":
    # 配置参数
    directory_path = r"C:\Users\Lenovo\Desktop\新建文件夹"  # 替换为你的目录路径
    output_filename = "合并文档.docx"
    
    merge_word_documents(directory_path, output_filename)
